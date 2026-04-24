from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document


@dataclass(frozen=True)
class RunData:
    text: str
    bold: bool
    strike: bool
    color: str | None


@dataclass(frozen=True)
class ParagraphData:
    text: str
    runs: list[RunData]


@dataclass(frozen=True)
class CellData:
    text: str
    paragraphs: list[ParagraphData]


@dataclass(frozen=True)
class RowData:
    cells: list[CellData]


@dataclass(frozen=True)
class TableData:
    rows: list[RowData]


@dataclass(frozen=True)
class DocumentData:
    paragraphs: list[ParagraphData]
    tables: list[TableData]
    headers: list[ParagraphData]
    footers: list[ParagraphData]


def read_docx(path: Path) -> DocumentData:
    document_path = Path(path)
    if not document_path.exists():
        raise FileNotFoundError(document_path)

    document = Document(str(document_path))
    headers: list[ParagraphData] = []
    footers: list[ParagraphData] = []

    for section in document.sections:
        headers.extend(_read_paragraphs(section.header.paragraphs))
        footers.extend(_read_paragraphs(section.footer.paragraphs))

    return DocumentData(
        paragraphs=_read_paragraphs(document.paragraphs),
        tables=_read_tables(document.tables),
        headers=headers,
        footers=footers,
    )


def _read_tables(tables) -> list[TableData]:
    return [
        TableData(
            rows=[
                RowData(
                    cells=[
                        CellData(
                            text=_normalize_text(cell.text),
                            paragraphs=_read_paragraphs(cell.paragraphs),
                        )
                        for cell in row.cells
                    ]
                )
                for row in table.rows
            ]
        )
        for table in tables
    ]


def _read_paragraphs(paragraphs) -> list[ParagraphData]:
    return [
        ParagraphData(
            text=paragraph.text,
            runs=[
                RunData(
                    text=run.text,
                    bold=bool(run.bold),
                    strike=bool(run.font.strike),
                    color=None if run.font.color.rgb is None else str(run.font.color.rgb),
                )
                for run in paragraph.runs
            ],
        )
        for paragraph in paragraphs
    ]


def _normalize_text(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\n", " | ").strip()
