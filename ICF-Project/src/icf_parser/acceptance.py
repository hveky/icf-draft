from __future__ import annotations

from collections import Counter, defaultdict
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path
import re

from docx import Document

from src.icf_parser.delivery_model import DeliveryRow
from src.icf_parser.template_writer import REVISION_HEADER_MARKERS


GREEN_RGB = "70AD47"
BLACK_RGB_VALUES = {"000000", "00000000"}
ROW_MATCH_THRESHOLD = 0.45
GENERIC_SECTION_LABELS = {"正文"}
VERSION_TOPIC_KEYS = {"版本号", "版本号/日期", "版本号及日期"}
CENTER_METADATA_MARKERS = ("研究中心", "主要研究者")
OUTLINE_PREFIX_PATTERN = re.compile(r"^\s*\d+(?:\.\d+)*[.．、]?\s*")
PAGE_PATTERN = re.compile(r"P\s*(\d+)(?:\s*[-–]\s*P?\s*(\d+))?", re.IGNORECASE)
CHINESE_PAGE_PATTERN = re.compile(r"第\s*(\d+)\s*页")


@dataclass(frozen=True)
class AlignmentMetric:
    score: float
    matched: int
    total_expected: int
    total_actual: int


@dataclass(frozen=True)
class StyleAlignment:
    score: float
    strike_score: float
    highlight_score: float
    generated_strike_rows: int
    acceptance_strike_rows: int
    generated_highlight_rows: int
    acceptance_highlight_rows: int


@dataclass(frozen=True)
class AcceptanceRow:
    topic: str
    section_page: str
    old_text: str
    new_text: str
    reason: str
    old_has_strike: bool
    new_has_highlight: bool


@dataclass(frozen=True)
class AcceptanceDiffReport:
    row_count_generated: int
    row_count_acceptance: int
    version_row_first_match: bool
    topic_alignment: AlignmentMetric
    section_page_alignment: AlignmentMetric
    reason_alignment: AlignmentMetric
    table_trace_alignment: AlignmentMetric
    style_alignment: StyleAlignment
    missing_rows: list[dict[str, str]]
    unexpected_rows: list[dict[str, str]]
    warnings: list[str]


@dataclass(frozen=True)
class MatchCandidate:
    generated_index: int
    acceptance_index: int
    overall_score: float
    topic_score: float
    section_score: float
    reason_score: float
    text_score: float


@dataclass(frozen=True)
class FieldDiff:
    index: int
    field: str
    generated: str
    acceptance: str


@dataclass(frozen=True)
class OrderDiff:
    generated_index: int
    expected_index: int
    topic: str
    section_page: str


@dataclass(frozen=True)
class StyleDiff:
    index: int
    field: str
    generated: bool
    acceptance: bool


@dataclass(frozen=True)
class DeliveryValidationReport:
    compared: bool
    delivery_passed: bool
    row_count_generated: int
    row_count_acceptance: int
    field_diffs: list[FieldDiff]
    order_diffs: list[OrderDiff]
    style_diffs: list[StyleDiff]
    missing_rows: list[dict[str, str]]
    unexpected_rows: list[dict[str, str]]
    blocking_issues: list[str]
    diagnostic_alignment: AcceptanceDiffReport | None


def compare_generated_to_acceptance(generated_docx: Path, acceptance_docx: Path) -> AcceptanceDiffReport:
    generated_rows = extract_delivery_rows_from_docx(Path(generated_docx))
    acceptance_rows = extract_delivery_rows_from_docx(Path(acceptance_docx))
    return compare_delivery_rows_to_acceptance_rows(generated_rows, acceptance_rows)


def compare_delivery_rows_to_acceptance_rows(
    generated_rows: list[DeliveryRow],
    acceptance_rows: list[DeliveryRow],
) -> AcceptanceDiffReport:
    generated_acceptance_rows = [_acceptance_row_from_delivery_row(row) for row in generated_rows]
    acceptance_acceptance_rows = [_acceptance_row_from_delivery_row(row) for row in acceptance_rows]
    matches = _match_rows(generated_acceptance_rows, acceptance_acceptance_rows)

    topic_alignment = _build_alignment_metric(generated_acceptance_rows, acceptance_acceptance_rows, matches, "topic")
    section_alignment = _build_alignment_metric(generated_acceptance_rows, acceptance_acceptance_rows, matches, "section_page")
    reason_alignment = _build_alignment_metric(generated_acceptance_rows, acceptance_acceptance_rows, matches, "reason")
    table_trace_alignment = _build_table_trace_metric(generated_acceptance_rows, acceptance_acceptance_rows, matches)
    style_alignment = _build_style_alignment(generated_acceptance_rows, acceptance_acceptance_rows, matches)

    missing_rows = [
        _row_snapshot(acceptance_acceptance_rows[acceptance_index])
        for acceptance_index in range(len(acceptance_acceptance_rows))
        if acceptance_index not in {acceptance_index for _, acceptance_index, _ in matches}
    ]
    unexpected_rows = [
        _row_snapshot(generated_acceptance_rows[generated_index])
        for generated_index in range(len(generated_acceptance_rows))
        if generated_index not in {generated_index for generated_index, _, _ in matches}
    ]

    version_row_first_match = bool(
        generated_acceptance_rows
        and _is_version_row(generated_acceptance_rows[0])
        and any(_is_version_row(row) for row in acceptance_acceptance_rows)
    )

    warnings = _build_warnings(
        generated_acceptance_rows,
        acceptance_acceptance_rows,
        topic_alignment,
        section_alignment,
        reason_alignment,
        table_trace_alignment,
        style_alignment,
        missing_rows,
        unexpected_rows,
        version_row_first_match,
    )

    return AcceptanceDiffReport(
        row_count_generated=len(generated_acceptance_rows),
        row_count_acceptance=len(acceptance_acceptance_rows),
        version_row_first_match=version_row_first_match,
        topic_alignment=topic_alignment,
        section_page_alignment=section_alignment,
        reason_alignment=reason_alignment,
        table_trace_alignment=table_trace_alignment,
        style_alignment=style_alignment,
        missing_rows=missing_rows,
        unexpected_rows=unexpected_rows,
        warnings=warnings,
    )


def extract_delivery_rows_from_docx(path: Path) -> list[DeliveryRow]:
    document = Document(str(path))
    table = _find_revision_table(document)
    rows: list[DeliveryRow] = []
    for row in table.rows[1:]:
        if len(row.cells) < 5:
            continue
        topic = row.cells[0].text.strip()
        section_page = row.cells[1].text.strip()
        old_text = row.cells[2].text.strip()
        new_text = row.cells[3].text.strip()
        reason = row.cells[4].text.strip()
        if not any((topic, section_page, old_text, new_text, reason)):
            continue
        rows.append(
            DeliveryRow(
                topic=topic,
                section_page=section_page,
                old_text=old_text,
                new_text=new_text,
                reason=reason,
                old_has_strike=_cell_has_strike(row.cells[2]),
                new_has_highlight=_cell_has_highlight(row.cells[3]),
                origin="document",
            )
        )
    return rows


def validate_generated_docx_against_acceptance(
    generated_docx: Path,
    acceptance_docx: Path,
) -> DeliveryValidationReport:
    generated_rows = extract_delivery_rows_from_docx(Path(generated_docx))
    acceptance_rows = extract_delivery_rows_from_docx(Path(acceptance_docx))
    diagnostic_alignment = compare_delivery_rows_to_acceptance_rows(generated_rows, acceptance_rows)
    return validate_delivery_rows(generated_rows, acceptance_rows, diagnostic_alignment=diagnostic_alignment)


def validate_delivery_rows(
    generated_rows: list[DeliveryRow],
    acceptance_rows: list[DeliveryRow],
    *,
    diagnostic_alignment: AcceptanceDiffReport | None = None,
) -> DeliveryValidationReport:
    if diagnostic_alignment is None:
        diagnostic_alignment = compare_delivery_rows_to_acceptance_rows(generated_rows, acceptance_rows)

    generated_nonstyle = [_non_style_fingerprint(row) for row in generated_rows]
    acceptance_nonstyle = [_non_style_fingerprint(row) for row in acceptance_rows]

    acceptance_positions: dict[tuple[str, str, str, str, str], list[int]] = defaultdict(list)
    for index, fingerprint in enumerate(acceptance_nonstyle):
        acceptance_positions[fingerprint].append(index)

    field_diffs: list[FieldDiff] = []
    order_diffs: list[OrderDiff] = []
    style_diffs: list[StyleDiff] = []

    for index in range(min(len(generated_rows), len(acceptance_rows))):
        generated_row = generated_rows[index]
        acceptance_row = acceptance_rows[index]
        generated_key = generated_nonstyle[index]
        acceptance_key = acceptance_nonstyle[index]

        if generated_key != acceptance_key:
            alternative_indexes = [candidate for candidate in acceptance_positions.get(generated_key, []) if candidate != index]
            if alternative_indexes:
                order_diffs.append(
                    OrderDiff(
                        generated_index=index,
                        expected_index=alternative_indexes[0],
                        topic=generated_row.topic,
                        section_page=generated_row.section_page,
                    )
                )
                continue
            field_diffs.extend(_build_field_diffs(index, generated_row, acceptance_row))
            continue

        if generated_row.old_has_strike != acceptance_row.old_has_strike:
            style_diffs.append(
                StyleDiff(
                    index=index,
                    field="old_has_strike",
                    generated=generated_row.old_has_strike,
                    acceptance=acceptance_row.old_has_strike,
                )
            )
        if generated_row.new_has_highlight != acceptance_row.new_has_highlight:
            style_diffs.append(
                StyleDiff(
                    index=index,
                    field="new_has_highlight",
                    generated=generated_row.new_has_highlight,
                    acceptance=acceptance_row.new_has_highlight,
                )
            )

    missing_rows = _counted_row_differences(acceptance_rows, Counter(acceptance_nonstyle) - Counter(generated_nonstyle))
    unexpected_rows = _counted_row_differences(generated_rows, Counter(generated_nonstyle) - Counter(acceptance_nonstyle))

    blocking_issues: list[str] = []
    if len(generated_rows) != len(acceptance_rows):
        blocking_issues.append("规范化交付行数与验收文件不一致。")
    if field_diffs:
        blocking_issues.append(f"存在 {len(field_diffs)} 处关键字段差异。")
    if order_diffs:
        blocking_issues.append(f"存在 {len(order_diffs)} 处行顺序差异。")
    if style_diffs:
        blocking_issues.append(f"存在 {len(style_diffs)} 处关键样式差异。")
    if missing_rows:
        blocking_issues.append(f"存在 {len(missing_rows)} 条验收行缺失。")
    if unexpected_rows:
        blocking_issues.append(f"存在 {len(unexpected_rows)} 条非验收生成行。")

    return DeliveryValidationReport(
        compared=True,
        delivery_passed=not blocking_issues,
        row_count_generated=len(generated_rows),
        row_count_acceptance=len(acceptance_rows),
        field_diffs=field_diffs,
        order_diffs=order_diffs,
        style_diffs=style_diffs,
        missing_rows=missing_rows,
        unexpected_rows=unexpected_rows,
        blocking_issues=blocking_issues,
        diagnostic_alignment=diagnostic_alignment,
    )


def _counted_row_differences(
    rows: list[DeliveryRow],
    counter: Counter[tuple[str, str, str, str, str]],
) -> list[dict[str, str]]:
    remaining = dict(counter)
    snapshots: list[dict[str, str]] = []
    for row in rows:
        key = _non_style_fingerprint(row)
        count = remaining.get(key, 0)
        if count <= 0:
            continue
        snapshots.append(_delivery_row_snapshot(row))
        remaining[key] = count - 1
    return snapshots


def _build_field_diffs(index: int, generated_row: DeliveryRow, acceptance_row: DeliveryRow) -> list[FieldDiff]:
    generated_values = _normalized_field_values(generated_row)
    acceptance_values = _normalized_field_values(acceptance_row)
    diffs: list[FieldDiff] = []
    for field_name in ("topic", "section_page", "old_text", "new_text", "reason"):
        if generated_values[field_name] == acceptance_values[field_name]:
            continue
        diffs.append(
            FieldDiff(
                index=index,
                field=field_name,
                generated=getattr(generated_row, field_name),
                acceptance=getattr(acceptance_row, field_name),
            )
        )
    return diffs


def _normalized_field_values(row: DeliveryRow) -> dict[str, str]:
    return {
        "topic": _normalized_topic_value(row),
        "section_page": _normalized_section_page_value(row),
        "old_text": _normalize_cell_text(row.old_text),
        "new_text": _normalize_cell_text(row.new_text),
        "reason": _normalize_reason_key(row.reason),
    }


def _non_style_fingerprint(row: DeliveryRow) -> tuple[str, str, str, str, str]:
    values = _normalized_field_values(row)
    return (
        values["topic"],
        values["section_page"],
        values["old_text"],
        values["new_text"],
        values["reason"],
    )


def _normalized_topic_value(row: DeliveryRow) -> str:
    acceptance_row = _acceptance_row_from_delivery_row(row)
    if _is_center_metadata_row(acceptance_row):
        return "中心信息"
    return _normalize_topic_key(row.topic)


def _normalized_section_page_value(row: DeliveryRow) -> str:
    acceptance_row = _acceptance_row_from_delivery_row(row)
    if _is_version_row(acceptance_row):
        return "页脚"
    if _is_center_metadata_row(acceptance_row):
        return "中心信息"
    label = _section_label_for_similarity(acceptance_row)
    pages = _extract_page_numbers(acceptance_row.section_page)
    if not pages:
        return label
    return f"{label}|{_compress_pages(pages)}"


def _compress_pages(pages: list[int]) -> str:
    if not pages:
        return ""
    ranges: list[str] = []
    start = pages[0]
    end = pages[0]
    for page in pages[1:]:
        if page == end + 1:
            end = page
            continue
        ranges.append(f"P{start}" if start == end else f"P{start}-P{end}")
        start = page
        end = page
    ranges.append(f"P{start}" if start == end else f"P{start}-P{end}")
    return ",".join(ranges)


def _delivery_row_snapshot(row: DeliveryRow) -> dict[str, str]:
    return {
        "topic": row.topic,
        "section_page": row.section_page,
        "reason": row.reason,
        "old_text_preview": row.old_text[:80],
        "new_text_preview": row.new_text[:80],
    }


def _normalize_cell_text(text: str) -> str:
    lines = [segment.strip() for segment in text.splitlines() if segment.strip()]
    if not lines:
        return ""
    return "\n".join(_normalize_text(line) for line in lines)


def _acceptance_row_from_delivery_row(row: DeliveryRow) -> AcceptanceRow:
    return AcceptanceRow(
        topic=row.topic,
        section_page=row.section_page,
        old_text=row.old_text,
        new_text=row.new_text,
        reason=row.reason,
        old_has_strike=row.old_has_strike,
        new_has_highlight=row.new_has_highlight,
    )


def _find_revision_table(document: Document):
    for table in document.tables:
        for row in table.rows:
            normalized = [_normalize_text(cell.text) for cell in row.cells]
            if len(normalized) < len(REVISION_HEADER_MARKERS):
                continue
            if all(marker in normalized[index] for index, marker in enumerate(REVISION_HEADER_MARKERS)):
                return table
    raise ValueError("revision table not found")


def _cell_has_strike(cell) -> bool:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.text and bool(run.font.strike):
                return True
    return False


def _cell_has_highlight(cell) -> bool:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            color = None
            if run.font.color is not None and run.font.color.rgb is not None:
                color = str(run.font.color.rgb)
            if run.text and bool(run.bold) and (color == GREEN_RGB or color not in BLACK_RGB_VALUES):
                return True
    return False


def _match_rows(generated_rows: list[AcceptanceRow], acceptance_rows: list[AcceptanceRow]) -> list[tuple[int, int, float]]:
    ranked_candidates: dict[int, list[MatchCandidate]] = {index: [] for index in range(len(generated_rows))}
    for generated_index, generated_row in enumerate(generated_rows):
        for acceptance_index, acceptance_row in enumerate(acceptance_rows):
            candidate = _build_match_candidate(generated_index, generated_row, acceptance_index, acceptance_row)
            if candidate is not None:
                ranked_candidates[generated_index].append(candidate)

    for generated_index in ranked_candidates:
        ranked_candidates[generated_index].sort(key=_match_preference_key, reverse=True)

    accepted: dict[int, MatchCandidate] = {}
    next_choice_index = {generated_index: 0 for generated_index in ranked_candidates}
    free_generated = list(range(len(generated_rows)))

    while free_generated:
        generated_index = free_generated.pop(0)
        candidates = ranked_candidates[generated_index]
        while next_choice_index[generated_index] < len(candidates):
            candidate = candidates[next_choice_index[generated_index]]
            next_choice_index[generated_index] += 1
            current = accepted.get(candidate.acceptance_index)
            if current is None:
                accepted[candidate.acceptance_index] = candidate
                break
            if _match_preference_key(candidate) > _match_preference_key(current):
                accepted[candidate.acceptance_index] = candidate
                free_generated.append(current.generated_index)
                break

    return sorted((candidate.generated_index, candidate.acceptance_index, candidate.overall_score) for candidate in accepted.values())


def _row_match_score(generated_row: AcceptanceRow, acceptance_row: AcceptanceRow) -> float:
    topic_score = _topic_similarity(generated_row, acceptance_row)
    section_score = _section_page_similarity(generated_row, acceptance_row)
    reason_score = _reason_similarity(generated_row.reason, acceptance_row.reason)
    old_text_score = _text_similarity(generated_row.old_text, acceptance_row.old_text)
    new_text_score = _text_similarity(generated_row.new_text, acceptance_row.new_text)
    return round(
        topic_score * 0.40
        + section_score * 0.25
        + reason_score * 0.20
        + old_text_score * 0.075
        + new_text_score * 0.075,
        6,
    )


def _build_alignment_metric(
    generated_rows: list[AcceptanceRow],
    acceptance_rows: list[AcceptanceRow],
    matches: list[tuple[int, int, float]],
    field_name: str,
) -> AlignmentMetric:
    if not acceptance_rows and not generated_rows:
        return AlignmentMetric(score=1.0, matched=0, total_expected=0, total_actual=0)
    if not matches:
        return AlignmentMetric(score=0.0, matched=0, total_expected=len(acceptance_rows), total_actual=len(generated_rows))

    score = sum(
        _field_similarity(generated_rows[generated_index], acceptance_rows[acceptance_index], field_name)
        for generated_index, acceptance_index, _ in matches
    ) / len(matches)

    return AlignmentMetric(
        score=round(score, 4),
        matched=len(matches),
        total_expected=len(acceptance_rows),
        total_actual=len(generated_rows),
    )


def _build_table_trace_metric(
    generated_rows: list[AcceptanceRow],
    acceptance_rows: list[AcceptanceRow],
    matches: list[tuple[int, int, float]],
) -> AlignmentMetric:
    if not acceptance_rows and not generated_rows:
        return AlignmentMetric(score=1.0, matched=0, total_expected=0, total_actual=0)
    if not matches:
        return AlignmentMetric(score=0.0, matched=0, total_expected=len(acceptance_rows), total_actual=len(generated_rows))

    score = sum(
        int(_has_table_trace(generated_rows[generated_index]) == _has_table_trace(acceptance_rows[acceptance_index]))
        for generated_index, acceptance_index, _ in matches
    ) / len(matches)

    return AlignmentMetric(
        score=round(score, 4),
        matched=len(matches),
        total_expected=len(acceptance_rows),
        total_actual=len(generated_rows),
    )


def _build_style_alignment(
    generated_rows: list[AcceptanceRow],
    acceptance_rows: list[AcceptanceRow],
    matches: list[tuple[int, int, float]],
) -> StyleAlignment:
    if not acceptance_rows and not generated_rows:
        return StyleAlignment(
            score=1.0,
            strike_score=1.0,
            highlight_score=1.0,
            generated_strike_rows=0,
            acceptance_strike_rows=0,
            generated_highlight_rows=0,
            acceptance_highlight_rows=0,
        )

    generated_strike_rows = sum(int(row.old_has_strike) for row in generated_rows)
    acceptance_strike_rows = sum(int(row.old_has_strike) for row in acceptance_rows)
    generated_highlight_rows = sum(int(row.new_has_highlight) for row in generated_rows)
    acceptance_highlight_rows = sum(int(row.new_has_highlight) for row in acceptance_rows)

    if not matches:
        return StyleAlignment(
            score=0.0,
            strike_score=0.0,
            highlight_score=0.0,
            generated_strike_rows=generated_strike_rows,
            acceptance_strike_rows=acceptance_strike_rows,
            generated_highlight_rows=generated_highlight_rows,
            acceptance_highlight_rows=acceptance_highlight_rows,
        )

    strike_score = sum(
        int(generated_rows[generated_index].old_has_strike == acceptance_rows[acceptance_index].old_has_strike)
        for generated_index, acceptance_index, _ in matches
    ) / len(matches)
    highlight_score = sum(
        int(generated_rows[generated_index].new_has_highlight == acceptance_rows[acceptance_index].new_has_highlight)
        for generated_index, acceptance_index, _ in matches
    ) / len(matches)

    return StyleAlignment(
        score=round((strike_score + highlight_score) / 2, 4),
        strike_score=round(strike_score, 4),
        highlight_score=round(highlight_score, 4),
        generated_strike_rows=generated_strike_rows,
        acceptance_strike_rows=acceptance_strike_rows,
        generated_highlight_rows=generated_highlight_rows,
        acceptance_highlight_rows=acceptance_highlight_rows,
    )


def _build_warnings(
    generated_rows: list[AcceptanceRow],
    acceptance_rows: list[AcceptanceRow],
    topic_alignment: AlignmentMetric,
    section_alignment: AlignmentMetric,
    reason_alignment: AlignmentMetric,
    table_trace_alignment: AlignmentMetric,
    style_alignment: StyleAlignment,
    missing_rows: list[dict[str, str]],
    unexpected_rows: list[dict[str, str]],
    version_row_first_match: bool,
) -> list[str]:
    warnings: list[str] = []
    if len(generated_rows) != len(acceptance_rows):
        warnings.append("生成结果与验收文件的修订记录行数不一致。")
    if not version_row_first_match:
        warnings.append("版本号/日期未在生成结果和验收文件中同时位于首行。")
    if topic_alignment.score < 0.8:
        warnings.append("主题对齐度偏低。")
    if section_alignment.score < 0.75:
        warnings.append("章节/页码对齐度偏低。")
    if reason_alignment.score < 0.75:
        warnings.append("更改原因措辞与验收文件存在明显差异。")
    if table_trace_alignment.score < 1.0:
        warnings.append("表格来源痕迹与验收文件未完全一致。")
    if style_alignment.score < 0.85:
        warnings.append("删除线或新增高亮样式与验收文件未完全一致。")
    if _has_center_summary(generated_rows) != _has_center_summary(acceptance_rows):
        warnings.append("中心信息合并结果与验收文件不一致。")
    if missing_rows:
        warnings.append(f"存在 {len(missing_rows)} 条验收行未在生成结果中匹配到。")
    if unexpected_rows:
        warnings.append(f"存在 {len(unexpected_rows)} 条生成行未在验收文件中匹配到。")
    return warnings


def _has_center_summary(rows: list[AcceptanceRow]) -> bool:
    return any(_is_center_metadata_row(row) for row in rows)


def _row_snapshot(row: AcceptanceRow) -> dict[str, str]:
    return {
        "topic": row.topic,
        "section_page": row.section_page,
        "reason": row.reason,
        "old_text_preview": row.old_text[:80],
        "new_text_preview": row.new_text[:80],
    }


def _has_table_trace(row: AcceptanceRow) -> bool:
    return "表格" in row.section_page or "#table" in row.section_page or row.section_page.startswith("表")


def _normalize_text(text: str) -> str:
    return "".join(text.split())


def _build_match_candidate(
    generated_index: int,
    generated_row: AcceptanceRow,
    acceptance_index: int,
    acceptance_row: AcceptanceRow,
) -> MatchCandidate | None:
    topic_score = _topic_similarity(generated_row, acceptance_row)
    section_score = _section_page_similarity(generated_row, acceptance_row)
    reason_score = _reason_similarity(generated_row.reason, acceptance_row.reason)
    text_score = (
        _text_similarity(generated_row.old_text, acceptance_row.old_text)
        + _text_similarity(generated_row.new_text, acceptance_row.new_text)
    ) / 2
    overall_score = _row_match_score(generated_row, acceptance_row)
    if overall_score < ROW_MATCH_THRESHOLD:
        return None
    return MatchCandidate(
        generated_index=generated_index,
        acceptance_index=acceptance_index,
        overall_score=overall_score,
        topic_score=topic_score,
        section_score=section_score,
        reason_score=reason_score,
        text_score=round(text_score, 6),
    )


def _match_preference_key(candidate: MatchCandidate) -> tuple[float, float, float, float, float, float]:
    return (
        candidate.topic_score,
        candidate.section_score,
        candidate.reason_score,
        candidate.text_score,
        candidate.overall_score,
        -abs(candidate.generated_index - candidate.acceptance_index),
    )


def _field_similarity(generated_row: AcceptanceRow, acceptance_row: AcceptanceRow, field_name: str) -> float:
    if field_name == "topic":
        return _topic_similarity(generated_row, acceptance_row)
    if field_name == "section_page":
        return _section_page_similarity(generated_row, acceptance_row)
    if field_name == "reason":
        return _reason_similarity(generated_row.reason, acceptance_row.reason)
    return _text_similarity(getattr(generated_row, field_name), getattr(acceptance_row, field_name))


def _topic_similarity(generated_row: AcceptanceRow, acceptance_row: AcceptanceRow) -> float:
    if _is_version_row(generated_row) and _is_version_row(acceptance_row):
        return 1.0
    if _is_center_metadata_row(generated_row) and _is_center_metadata_row(acceptance_row):
        return 1.0
    return _text_similarity(_normalize_topic_key(generated_row.topic), _normalize_topic_key(acceptance_row.topic))


def _section_page_similarity(generated_row: AcceptanceRow, acceptance_row: AcceptanceRow) -> float:
    if _is_version_row(generated_row) and _is_version_row(acceptance_row):
        return 1.0
    if _is_center_metadata_row(generated_row) and _is_center_metadata_row(acceptance_row):
        return 1.0

    label_score = _text_similarity(
        _section_label_for_similarity(generated_row),
        _section_label_for_similarity(acceptance_row),
    )
    page_score = _page_similarity(
        _extract_page_numbers(generated_row.section_page),
        _extract_page_numbers(acceptance_row.section_page),
    )
    return round(label_score * 0.75 + page_score * 0.25, 6)


def _reason_similarity(left: str, right: str) -> float:
    return _text_similarity(_normalize_reason_key(left), _normalize_reason_key(right))


def _normalize_topic_key(text: str) -> str:
    normalized = _normalize_text(_strip_outline_prefix(text))
    if normalized in {"版本号及日期", "版本号/日期修订", "版本更新"}:
        return "版本号/日期"
    return normalized.replace("版本号及日期", "版本号/日期")


def _normalize_reason_key(text: str) -> str:
    normalized = _normalize_text(text)
    normalized = normalized.rstrip("。；;")
    if "版本" in normalized and ("更新" in normalized or "修订" in normalized):
        return "版本更新"
    return normalized


def _strip_outline_prefix(text: str) -> str:
    stripped_lines: list[str] = []
    for raw_line in text.splitlines() or [text]:
        line = OUTLINE_PREFIX_PATTERN.sub("", raw_line.strip())
        stripped_lines.append(line)
    return "\n".join(stripped_lines)


def _section_label_for_similarity(row: AcceptanceRow) -> str:
    if _is_version_row(row):
        return "页脚"
    if _is_center_metadata_row(row):
        return "中心信息"

    labels: list[str] = []
    for raw_line in row.section_page.splitlines() or [row.section_page]:
        line = PAGE_PATTERN.sub("", _strip_outline_prefix(raw_line))
        line = CHINESE_PAGE_PATTERN.sub("", line)
        line = line.replace("：", "").replace(":", "").strip(" /-")
        normalized = _normalize_text(line)
        if normalized:
            labels.append(normalized)

    label = "\n".join(labels)
    if not label or label in GENERIC_SECTION_LABELS:
        return _normalize_topic_key(row.topic)
    return label


def _extract_page_numbers(text: str) -> list[int]:
    pages: list[int] = []
    for start_text, end_text in PAGE_PATTERN.findall(text):
        start = int(start_text)
        end = int(end_text) if end_text else start
        if end < start:
            start, end = end, start
        pages.extend(range(start, end + 1))
    for page_text in CHINESE_PAGE_PATTERN.findall(text):
        pages.append(int(page_text))
    return sorted(set(pages))


def _page_similarity(left_pages: list[int], right_pages: list[int]) -> float:
    if not left_pages and not right_pages:
        return 1.0
    if not left_pages or not right_pages:
        return 0.6

    left_set = set(left_pages)
    right_set = set(right_pages)
    overlap = len(left_set & right_set)
    if overlap:
        return round(overlap / len(left_set | right_set), 6)

    distance = min(abs(left - right) for left in left_pages for right in right_pages)
    if distance <= 1:
        return 0.9
    if distance <= 2:
        return 0.75
    if distance <= 3:
        return 0.6
    return 0.25


def _is_version_row(row: AcceptanceRow) -> bool:
    return _normalize_topic_key(row.topic) in VERSION_TOPIC_KEYS


def _is_center_metadata_row(row: AcceptanceRow) -> bool:
    if _normalize_topic_key(row.topic) == "中心信息":
        return True
    text = f"{row.old_text}\n{row.new_text}"
    if not any(marker in text for marker in CENTER_METADATA_MARKERS):
        return False
    return any(label in text for label in ("研究中心：", "研究中心:", "主要研究者：", "主要研究者:"))


def _text_similarity(left: str, right: str) -> float:
    if not left and not right:
        return 1.0
    return SequenceMatcher(a=_normalize_text(left), b=_normalize_text(right)).ratio()
